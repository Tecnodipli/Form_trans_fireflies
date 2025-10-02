# main.py
# -*- coding: utf-8 -*-
import os
import re
import zipfile
import io
import uuid
import json
import time
import httpx
import asyncio
from collections import OrderedDict, defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, BackgroundTasks
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ========================= 
# CONFIGURACIÓN FIREFLIES
# =========================
FIREFLIES_API_KEY = os.getenv("FIREFLIES_API_KEY", "tu_api_key_aqui")
FIREFLIES_GRAPHQL_URL = "https://api.fireflies.ai/graphql"

# ========================= 
# CONFIGURACIÓN URL PÚBLICA
# =========================
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "https://form-trans-fireflies.onrender.com")

# ========================= 
# 1) Configuración de formato
# ========================= 
FONT_NAME = "Arial"
FONT_SIZE_PT = 12
SPACE_AFTER_LABEL_PT = 12

# ========================= 
# 2) Patrones robustos
# ========================= 
TIME_ONLY = re.compile(r'^\s*\(?\d{1,2}:\d{2}(?::\d{2})?\)?\s*$')
INLINE_LABEL = re.compile(
    r'^\s*(?:\(?\d{1,2}:\d{2}(?::\d{2})?\)?\s*)?'
    r'(speaker\s*\d+)\s*:?\s*(\S.+)$',
    re.IGNORECASE
)
LABEL_ONLY = re.compile(
    r'^\s*(?:\(?\d{1,2}:\d{2}(?::\d{2})?\)?\s*)?'
    r'(speaker\s*\d+)\s*:?\s*$',
    re.IGNORECASE
)

# ========================= 
# 3) Utilidades de texto/docx
# ========================= 
def paragraph_text(p: Paragraph) -> str:
    return ''.join(r.text for r in p.runs) if p.runs else p.text

def normalize_label(lbl: str) -> str:
    return re.sub(r'\s+', ' ', (lbl or '')).strip().casefold()

def ensure_colon_upper(s: str) -> str:
    s = (s or '').strip()
    if not s.endswith(':'):
        s += ':'
    return s.upper()

def clear_paragraph(p: Paragraph):
    p.text = ''

def set_spacing(p: Paragraph, after_pt=SPACE_AFTER_LABEL_PT, before_pt=0):
    pf = p.paragraph_format
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)

def write_label_plus_content(
    p: Paragraph,
    final_label: str,
    content: str,
    bold_label: bool,
    bold_content: bool,
    apply_spacing: bool = True,
):
    content = re.sub(r'\s+', ' ', content or '').strip()
    clear_paragraph(p)
    r1 = p.add_run(final_label + ' ')
    r1.bold = bold_label
    r2 = p.add_run(content)
    r2.bold = bold_content
    if apply_spacing:
        set_spacing(p, after_pt=SPACE_AFTER_LABEL_PT)

def bold_whole_paragraph(p: Paragraph):
    if not p.runs:
        if p.text:
            txt = p.text
            clear_paragraph(p)
            r = p.add_run(txt)
            r.bold = True
        return
    for r in p.runs:
        r.bold = True

def is_time_only(p: Paragraph) -> bool:
    return bool(TIME_ONLY.match(paragraph_text(p).strip()))

def is_label_start(p: Paragraph):
    txt = paragraph_text(p)
    m = INLINE_LABEL.match(txt)
    if m:
        return ('inline', m.group(1), m.group(2))
    m = LABEL_ONLY.match(txt)
    if m:
        return ('only', m.group(1), None)
    return None

def apply_global_font(doc: Document, name=FONT_NAME, size_pt=FONT_SIZE_PT):
    try:
        doc.styles['Normal'].font.name = name
        doc.styles['Normal'].font.size = Pt(size_pt)
    except Exception:
        pass
    
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = name
            if r._element.rPr is not None:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            r.font.size = Pt(size_pt)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = name
                        if r._element.rPr is not None:
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), name)
                        r.font.size = Pt(size_pt)

def fmt_hms(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

# ========================= 
# 4) Registro (TXT) y guardado seguro (en memoria)
# ========================= 
def write_txt_control_report_in_memory(log: dict, turn_logs: list):
    lines = []
    lines.append("REGISTRO DE CONTROL Y PROCESO")
    lines.append("=" * 34)
    lines.append(f"Fecha de procesamiento: {log['ts']}")
    lines.append(f"Archivo de entrada: {log['input_file']}")
    lines.append(f"Párrafos totales: {log['total_paragraphs']}")
    lines.append(f"Etiquetas detectadas: {', '.join(log['labels_detected']) if log['labels_detected'] else '—'}")
    lines.append(f"Párrafos actualizados: {log['changed_count']}")
    lines.append(f"Timestamps detectados: {log['time_only_count']}")
    lines.append(f"Duración TOTAL: {log['exec_total_hms']} ({log['exec_total_seconds']:.3f}s)")
    lines.append(f"Duración de PROCESAMIENTO: {log['exec_processing_hms']} ({log['exec_processing_seconds']:.3f}s)")
    lines.append("")
    lines.append("Mapeo aplicado (detectada -> final | turnos):")
    for k, raw in log['mapping_raw_order']:
        final = log['mapping'][k]
        cnt = log['counts_by_final'][final]
        lines.append(f"- {raw} -> {final} | {cnt}")
    lines.append("")
    lines.append("Detalle por turno:")
    lines.append("index\traw_label\tfinal_label\tcase\tcontent_found\tinterviewer\tstart_par\tend_par")
    for row in turn_logs:
        lines.append(f"{row['index']}\t{row['raw_label']}\t{row['final_label']}\t{row['kind']}\t{row['content_found']}\t{row['interviewer']}\t{row['start_par']}\t{row['end_par']}")
    
    return "\n".join(lines)

# ========================= 
# 5) Helpers de roles
# ========================= 
def normalize_role_label(label: str) -> str:
    s = re.sub(r'\s+', ' ', (label or '')).strip()
    if s.endswith(':'):
        s = s[:-1]
    return s.casefold()

def is_interviewer_final(label: str) -> bool:
    norm = normalize_role_label(label)
    return 'entrevistador' in norm

# ========================= 
# 6) FUNCIONES FIREFLIES
# ========================= 

# Directorio temporal para archivos
TEMP_UPLOADS_DIR = Path("temp_uploads")
TEMP_UPLOADS_DIR.mkdir(exist_ok=True)

async def upload_audio_to_fireflies_via_url(audio_url: str, title: str) -> str:
    """
    Sube un archivo de audio/video a Fireflies mediante URL pública
    """
    mutation = """
    mutation UploadAudio($input: UploadAudioInput!) {
      uploadAudio(input: $input) {
        success
        title
        message
      }
    }
    """
    
    headers = {
        'Authorization': f'Bearer {FIREFLIES_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            FIREFLIES_GRAPHQL_URL,
            headers=headers,
            json={
                'query': mutation,
                'variables': {
                    'input': {
                        'url': audio_url,
                        'title': title
                    }
                }
            }
        )
        
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Error al subir a Fireflies: {response.text}")
        
        result = response.json()
        
        if 'errors' in result:
            raise HTTPException(status_code=500, detail=f"Error de Fireflies: {result['errors']}")
        
        # Fireflies no retorna transcript_id inmediatamente, hay que buscarlo después
        return title

async def find_transcript_by_title(title: str, max_attempts: int = 20) -> Optional[str]:
    """
    Busca una transcripción por título (puede tardar unos segundos en aparecer)
    """
    query = """
    query Transcripts {
      transcripts(limit: 50) {
        id
        title
        date
      }
    }
    """
    
    headers = {
        'Authorization': f'Bearer {FIREFLIES_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    for attempt in range(max_attempts):
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                FIREFLIES_GRAPHQL_URL,
                headers=headers,
                json={'query': query}
            )
            
            if response.status_code == 200:
                result = response.json()
                
                if 'data' in result and 'transcripts' in result['data']:
                    for transcript in result['data']['transcripts']:
                        if transcript['title'] == title:
                            return transcript['id']
        
        # Esperar 3 segundos antes de reintentar
        await asyncio.sleep(3)
    
    return None

async def check_transcription_status(transcript_id: str) -> dict:
    """
    Verifica el estado de la transcripción
    """
    query = """
    query Transcript($transcriptId: String!) {
      transcript(id: $transcriptId) {
        id
        title
        transcript_url
        audio_url
        date
        duration
        sentences {
          text
          speaker_name
          start_time
          end_time
        }
      }
    }
    """
    
    headers = {
        'Authorization': f'Bearer {FIREFLIES_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            FIREFLIES_GRAPHQL_URL,
            headers=headers,
            json={
                'query': query,
                'variables': {'transcriptId': transcript_id}
            }
        )
        
        if response.status_code != 200:
            return {'status': 'error', 'message': response.text}
        
        result = response.json()
        
        if 'errors' in result:
            return {'status': 'error', 'message': result['errors']}
        
        transcript_data = result['data']['transcript']
        
        if transcript_data and transcript_data.get('sentences'):
            return {'status': 'completed', 'data': transcript_data}
        else:
            return {'status': 'processing'}

def create_docx_from_fireflies(transcript_data: dict) -> io.BytesIO:
    """
    Crea un documento .docx a partir de la transcripción de Fireflies
    """
    doc = Document()
    
    current_speaker = None
    current_text = []
    
    for sentence in transcript_data['sentences']:
        speaker = sentence['speaker_name'] or "Speaker Unknown"
        text = sentence['text']
        
        if speaker != current_speaker:
            if current_speaker and current_text:
                p = doc.add_paragraph()
                p.add_run(f"{current_speaker}:").bold = True
                p.add_run(" " + " ".join(current_text))
            
            current_speaker = speaker
            current_text = [text]
        else:
            current_text.append(text)
    
    if current_speaker and current_text:
        p = doc.add_paragraph()
        p.add_run(f"{current_speaker}:").bold = True
        p.add_run(" " + " ".join(current_text))
    
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream

def save_temp_file(file_content: bytes, filename: str) -> tuple[str, str]:
    """
    Guarda el archivo temporalmente y retorna el path y token
    """
    file_token = str(uuid.uuid4())
    extension = Path(filename).suffix
    temp_filename = f"{file_token}{extension}"
    temp_path = TEMP_UPLOADS_DIR / temp_filename
    
    with open(temp_path, 'wb') as f:
        f.write(file_content)
    
    return str(temp_path), file_token

def cleanup_temp_file(file_path: str):
    """
    Elimina el archivo temporal
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        print(f"Error al eliminar archivo temporal: {e}")

# ========================= 
# 7) Procesamiento principal (versión API)
# ========================= 
def process_file_api(file_stream: io.BytesIO, interview_type: str, label_mapping_user: dict = None, file_name: str = "file.docx"):
    t0_total = time.perf_counter()
    
    doc = Document(file_stream)
    
    found_labels = OrderedDict()
    for p in doc.paragraphs:
        hit = is_label_start(p)
        if hit:
            _, raw_label, _ = hit
            k = normalize_label(raw_label)
            if k not in found_labels:
                found_labels[k] = raw_label.strip()
    
    label_mapping = {}
    if label_mapping_user:
        for k, raw in found_labels.items():
            if k in label_mapping_user:
                final_label = ensure_colon_upper(label_mapping_user[k] or raw)
            else:
                final_label = ensure_colon_upper(raw)
            label_mapping[k] = final_label
    else:
        for k, raw in found_labels.items():
            label_mapping[k] = ensure_colon_upper(raw)
    
    t0_processing = time.perf_counter()
    
    i = 0
    changed = 0
    n = len(doc.paragraphs)
    time_only_count = sum(1 for p in doc.paragraphs if is_time_only(p))
    
    counts_by_final = defaultdict(int)
    turn_logs = []
    
    while i < n:
        p = doc.paragraphs[i]
        hit = is_label_start(p)
        
        if not hit:
            i += 1
            continue
        
        kind, raw_label, content_inline = hit
        key = normalize_label(raw_label)
        final_label = label_mapping.get(key)
        
        if not final_label:
            i += 1
            continue
        
        is_interviewer = is_interviewer_final(final_label)
        bold_label_flag = is_interviewer
        bold_content_flag = is_interviewer
        
        start_par = i
        content_found = False
        end_par = i
        
        if kind == 'inline':
            write_label_plus_content(p, final_label, content_inline, bold_label_flag, bold_content_flag, apply_spacing=True)
            content_found = True
            changed += 1
            counts_by_final[final_label] += 1
            turn_logs.append({
                'index': len(turn_logs) + 1,
                'raw_label': raw_label,
                'final_label': final_label,
                'kind': 'inline',
                'content_found': content_found,
                'interviewer': is_interviewer,
                'start_par': start_par,
                'end_par': end_par
            })
            i += 1
            continue
        
        j = i + 1
        while j < n:
            txtj = paragraph_text(doc.paragraphs[j]).strip()
            if not txtj or is_time_only(doc.paragraphs[j]):
                j += 1
                continue
            if is_label_start(doc.paragraphs[j]):
                break
            break
        
        if j >= n or is_label_start(doc.paragraphs[j]) or not paragraph_text(doc.paragraphs[j]).strip() or is_time_only(doc.paragraphs[j]):
            write_label_plus_content(p, final_label, "", bold_label_flag, bold_content_flag, apply_spacing=True)
            changed += 1
            counts_by_final[final_label] += 1
            turn_logs.append({
                'index': len(turn_logs) + 1,
                'raw_label': raw_label,
                'final_label': final_label,
                'kind': 'only',
                'content_found': False,
                'interviewer': is_interviewer,
                'start_par': start_par,
                'end_par': start_par
            })
            i += 1
            continue
        
        first_content = paragraph_text(doc.paragraphs[j])
        write_label_plus_content(p, final_label, first_content, bold_label_flag, bold_content_flag, apply_spacing=True)
        clear_paragraph(doc.paragraphs[j])
        content_found = True
        changed += 1
        counts_by_final[final_label] += 1
        
        k = j + 1
        last_non_time_idx = i
        while k < n and not is_label_start(doc.paragraphs[k]):
            if not is_time_only(doc.paragraphs[k]) and paragraph_text(doc.paragraphs[k]).strip():
                last_non_time_idx = k
                if is_interviewer:
                    bold_whole_paragraph(doc.paragraphs[k])
            k += 1
        
        set_spacing(doc.paragraphs[last_non_time_idx], after_pt=SPACE_AFTER_LABEL_PT)
        end_par = last_non_time_idx
        
        turn_logs.append({
            'index': len(turn_logs) + 1,
            'raw_label': raw_label,
            'final_label': final_label,
            'kind': 'only+merge',
            'content_found': content_found,
            'interviewer': is_interviewer,
            'start_par': start_par,
            'end_par': end_par
        })
        
        i = k
    
    apply_global_font(doc, name=FONT_NAME, size_pt=FONT_SIZE_PT)
    
    t1_total = time.perf_counter()
    exec_total = t1_total - t0_total
    exec_processing = t1_total - t0_processing
    
    log = {
        'ts': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'input_file': file_name,
        'total_paragraphs': len(doc.paragraphs),
        'labels_detected': list(found_labels.values()),
        'mapping': label_mapping,
        'mapping_raw_order': list(found_labels.items()),
        'changed_count': changed,
        'time_only_count': time_only_count,
        'counts_by_final': counts_by_final,
        'exec_total_seconds': exec_total,
        'exec_total_hms': fmt_hms(exec_total),
        'exec_processing_seconds': exec_processing,
        'exec_processing_hms': fmt_hms(exec_processing),
    }
    
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    txt_content = write_txt_control_report_in_memory(log, turn_logs)
    txt_stream = io.BytesIO(txt_content.encode('utf-8'))
    txt_stream.seek(0)
    
    return docx_stream, txt_stream

def detect_labels_api(file_stream: io.BytesIO):
    doc = Document(file_stream)
    
    found_labels = OrderedDict()
    for p in doc.paragraphs:
        hit = is_label_start(p)
        if hit:
            _, raw_label, _ = hit
            k = normalize_label(raw_label)
            if k not in found_labels:
                found_labels[k] = raw_label.strip()
    
    return list(found_labels.values())

# ========================== 
# FastAPI
# ========================== 
app = FastAPI(title="Formateador de Transcripciones con Fireflies")

ALLOWED_ORIGINS = [
    "https://www.dipli.ai",
    "https://dipli.ai",
    "https://isagarcivill09.wixsite.com/turop",
    "https://isagarcivill09.wixsite.com/turop/tienda",
    "https://isagarcivill09-wixsite-com.filesusr.com",
    "https://www.dipli.ai/preparaci%C3%B3n",
    "https://www-dipli-ai.filesusr.com"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Almacenamiento temporal
DOWNLOADS = {}
TRANSCRIPTION_JOBS = {}
EXP_MINUTES = 30

def cleanup_downloads():
    now = datetime.utcnow()
    expired = [t for t, (_, exp) in DOWNLOADS.items() if exp <= now]
    for t in expired:
        DOWNLOADS.pop(t, None)

# ========================== 
# ENDPOINTS FIREFLIES (FLUJO COMPLETO)
# ========================== 

@app.post("/fireflies/subir_audio/")
async def subir_audio_fireflies(
    file: UploadFile = File(...),
    interview_type: str = Form(...)
):
    """
    Paso 1: Sube un archivo de audio/video a Fireflies para transcripción.
    Retorna un job_id para consultar el estado.
    """
    temp_path = None
    try:
        if not (file.filename.endswith('.mp3') or file.filename.endswith('.mp4')):
            raise HTTPException(status_code=400, detail="El archivo debe ser MP3 o MP4")
        
        file_content = await file.read()
        
        if len(file_content) > 1.5 * 1024 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="El archivo no debe superar 1.5 GB")
        
        # Guardar archivo temporalmente
        temp_path, file_token = save_temp_file(file_content, file.filename)
        extension = Path(file.filename).suffix
        
        # Construir URL pública usando la configuración
        public_url = f"{PUBLIC_BASE_URL}/temp_audio/{file_token}{extension}"
        
        print(f"URL pública generada: {public_url}")
        
        # Subir a Fireflies mediante URL
        title = f"{interview_type} - {file.filename} - {file_token}"
        upload_result = await upload_audio_to_fireflies_via_url(public_url, title)
        
        # Buscar el transcript_id (puede tardar unos segundos)
        print("Buscando transcript_id...")
        transcript_id = await find_transcript_by_title(title)
        
        if not transcript_id:
            raise HTTPException(status_code=500, detail="No se pudo obtener el transcript_id. Intenta de nuevo en unos segundos.")
        
        print(f"Transcript ID encontrado: {transcript_id}")
        
        # Crear job_id
        job_id = str(uuid.uuid4())
        
        # Guardar información del job
        TRANSCRIPTION_JOBS[job_id] = {
            'transcript_id': transcript_id,
            'interview_type': interview_type,
            'filename': file.filename,
            'status': 'processing',
            'created_at': datetime.utcnow(),
            'temp_path': temp_path,
            'file_token': file_token,
            'title': title
        }
        
        return JSONResponse(content={
            'job_id': job_id,
            'transcript_id': transcript_id,
            'status': 'processing',
            'message': 'Audio subido exitosamente. La transcripción está en proceso.',
            'public_url': public_url
        })
        
    except Exception as e:
        if temp_path:
            cleanup_temp_file(temp_path)
        raise HTTPException(status_code=500, detail=f"Error al subir audio: {str(e)}")

@app.get("/temp_audio/{file_token}")
async def serve_temp_audio(file_token: str):
    """
    Sirve archivos temporales para que Fireflies los descargue
    """
    # Buscar el archivo en el directorio temporal
    for temp_file in TEMP_UPLOADS_DIR.glob(f"{file_token}.*"):
        if temp_file.is_file():
            return StreamingResponse(
                open(temp_file, 'rb'),
                media_type='application/octet-stream'
            )
    
    raise HTTPException(status_code=404, detail="Archivo no encontrado")

@app.get("/fireflies/estado_transcripcion/{job_id}")
async def estado_transcripcion(job_id: str):
    """
    Paso 2: Consulta el estado de una transcripción en proceso.
    Llamar periódicamente hasta que status sea 'completed'.
    """
    if job_id not in TRANSCRIPTION_JOBS:
        raise HTTPException(status_code=404, detail="Job ID no encontrado")
    
    job = TRANSCRIPTION_JOBS[job_id]
    transcript_id = job['transcript_id']
    
    status_result = await check_transcription_status(transcript_id)
    
    if status_result['status'] == 'completed':
        job['status'] = 'completed'
        job['transcript_data'] = status_result['data']
        
        return JSONResponse(content={
            'job_id': job_id,
            'status': 'completed',
            'message': 'Transcripción completada. Puede proceder a detectar etiquetas.'
        })
    
    elif status_result['status'] == 'processing':
        return JSONResponse(content={
            'job_id': job_id,
            'status': 'processing',
            'message': 'La transcripción está en proceso. Por favor espere.'
        })
    
    else:
        job['status'] = 'error'
        return JSONResponse(content={
            'job_id': job_id,
            'status': 'error',
            'message': f"Error en la transcripción: {status_result.get('message', 'Error desconocido')}"
        })

@app.post("/fireflies/detectar_etiquetas/{job_id}")
async def detectar_etiquetas_fireflies(job_id: str):
    """
    Paso 3: Genera el documento .docx desde Fireflies y detecta etiquetas para mapeo.
    """
    if job_id not in TRANSCRIPTION_JOBS:
        raise HTTPException(status_code=404, detail="Job ID no encontrado")
    
    job = TRANSCRIPTION_JOBS[job_id]
    
    if job['status'] != 'completed':
        raise HTTPException(status_code=400, detail="La transcripción aún no está completa")
    
    try:
        docx_stream = create_docx_from_fireflies(job['transcript_data'])
        
        labels = detect_labels_api(docx_stream)
        
        docx_stream.seek(0)
        job['docx_content'] = docx_stream.read()
        
        if not labels:
            return JSONResponse(content={
                'job_id': job_id,
                'labels': [],
                'message': 'Transcripción generada pero no se detectaron etiquetas Speaker N.'
            })
        
        return JSONResponse(content={
            'job_id': job_id,
            'labels': labels,
            'message': 'Etiquetas detectadas exitosamente'
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al detectar etiquetas: {str(e)}")

@app.post("/fireflies/formatear/{job_id}")
async def formatear_desde_fireflies(
    job_id: str,
    label_mapping: str = Form("null"),
    background_tasks: BackgroundTasks = None
):
    """
    Paso 4: Formatea el documento generado por Fireflies con el mapeo seleccionado.
    """
    if job_id not in TRANSCRIPTION_JOBS:
        raise HTTPException(status_code=404, detail="Job ID no encontrado")
    
    job = TRANSCRIPTION_JOBS[job_id]
    
    if job['status'] != 'completed':
        raise HTTPException(status_code=400, detail="La transcripción aún no está completa")
    
    if 'docx_content' not in job:
        raise HTTPException(status_code=400, detail="Debe detectar etiquetas primero")
    
    try:
        file_stream = io.BytesIO(job['docx_content'])
        
        mapping_data = None
        if label_mapping and label_mapping != "null":
            try:
                mapping_data = json.loads(label_mapping)
            except json.JSONDecodeError:
                print("Advertencia: El 'label_mapping' no es un JSON válido.")
        
        if mapping_data:
            mapping_data_normalized = {normalize_label(k): v for k, v in mapping_data.items()}
        else:
            mapping_data_normalized = None
        
        docx_stream, txt_stream = process_file_api(
            file_stream=file_stream,
            interview_type=job['interview_type'],
            label_mapping_user=mapping_data_normalized,
            file_name=job['filename']
        )
        
        base_filename = Path(job['filename']).stem
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            docx_name = f"{base_filename}_formateado_FINAL.docx"
            zip_file.writestr(docx_name, docx_stream.getvalue())
            
            txt_name = f"{base_filename}_registro_control_proceso.txt"
            zip_file.writestr(txt_name, txt_stream.getvalue())
        
        zip_buffer.seek(0)
        
        token = str(uuid.uuid4())
        expiration = datetime.utcnow() + timedelta(minutes=EXP_MINUTES)
        DOWNLOADS[token] = (zip_buffer, expiration)
        
        # Limpiar archivo temporal en background
        if 'temp_path' in job and background_tasks:
            background_tasks.add_task(cleanup_temp_file, job['temp_path'])
        
        TRANSCRIPTION_JOBS.pop(job_id, None)
        
        return JSONResponse(content={
            'token': token,
            'filename': f"{base_filename}_formateado.zip",
            'message': 'Documento formateado exitosamente'
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al formatear: {str(e)}")

# ========================== 
# ENDPOINTS ORIGINALES (para subir .docx directamente)
# ========================== 

@app.post("/detectar_etiquetas/")
async def detectar_etiquetas(file: UploadFile = File(...)):
    """
    Detecta etiquetas "Speaker N" en un documento .docx ya transcrito.
    Usa POST para aceptar el archivo en el cuerpo de la solicitud.
    """
    try:
        file_content = await file.read()
        file_stream = io.BytesIO(file_content)
        
        labels = detect_labels_api(file_stream)
        
        if not labels:
            return JSONResponse(content={
                "labels": [],
                "message": "No se detectaron etiquetas 'Speaker N'."
            })
        
        return JSONResponse(content={"labels": labels})
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al detectar etiquetas: {e}")

@app.post("/formatear/")
async def formatear_transcripcion(
    file: UploadFile = File(...),
    interview_type: str = Form(...),
    label_mapping: str = Form("null")
):
    """
    Formatea un documento .docx ya transcrito y genera un archivo de registro.
    Acepta el archivo y el mapeo de etiquetas en el cuerpo de la solicitud.
    """
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="El archivo debe ser un .docx")
        
        file_content = await file.read()
        file_stream = io.BytesIO(file_content)
        
        mapping_data = None
        if label_mapping and label_mapping != "null":
            try:
                mapping_data = json.loads(label_mapping)
            except json.JSONDecodeError:
                print("Advertencia: El 'label_mapping' no es un JSON válido. Se usará el formato por defecto.")
        
        if mapping_data:
            mapping_data_normalized = {normalize_label(k): v for k, v in mapping_data.items()}
        else:
            mapping_data_normalized = None
        
        docx_stream, txt_stream = process_file_api(
            file_stream=file_stream,
            interview_type=interview_type,
            label_mapping_user=mapping_data_normalized,
            file_name=file.filename
        )
        
        base_filename = Path(file.filename).stem
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            docx_name = f"{base_filename}_formateado_FINAL.docx"
            zip_file.writestr(docx_name, docx_stream.getvalue())
            
            txt_name = f"{base_filename}_registro_control_proceso.txt"
            zip_file.writestr(txt_name, txt_stream.getvalue())
        
        zip_buffer.seek(0)
        
        token = str(uuid.uuid4())
        expiration = datetime.utcnow() + timedelta(minutes=EXP_MINUTES)
        DOWNLOADS[token] = (zip_buffer, expiration)
        
        return JSONResponse(content={
            "token": token,
            "filename": f"{base_filename}_formateado.zip"
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en el procesamiento: {e}")

@app.get("/descargar/{token}")
async def descargar_archivo(token: str):
    """
    Paso 5: Permite la descarga de un archivo comprimido por un token.
    Este endpoint sirve tanto para flujo Fireflies como para subida directa de .docx
    """
    cleanup_downloads()
    
    if token not in DOWNLOADS:
        raise HTTPException(status_code=404, detail="Token de descarga no válido o expirado.")
    
    zip_buffer, _ = DOWNLOADS.pop(token)
    
    response = StreamingResponse(
        io.BytesIO(zip_buffer.getvalue()),
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename=archivos_formateados.zip",
            "Content-Length": str(len(zip_buffer.getvalue()))
        }
    )
    
    return response

# ========================== 
# ENDPOINT DE SALUD
# ========================== 

@app.get("/")
async def root():
    """
    Endpoint de salud para verificar que la API está funcionando
    """
    return {
        "message": "API Formateador de Transcripciones con Fireflies",
        "status": "online",
        "version": "2.0",
        "endpoints": {
            "fireflies": [
                "POST /fireflies/subir_audio/",
                "GET /fireflies/estado_transcripcion/{job_id}",
                "POST /fireflies/detectar_etiquetas/{job_id}",
                "POST /fireflies/formatear/{job_id}"
            ],
            "direct": [
                "POST /detectar_etiquetas/",
                "POST /formatear/"
            ],
            "common": [
                "GET /descargar/{token}"
            ]
        }
    }
